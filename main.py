import tkinter as tk
from operator import index
from docx import Document
from tkinter import messagebox
import os
import taskcreator
###
import taskcreator as tc
from task_window import create_second_window

# this is OOP
def add_window():
    create_second_window(root)
def create_word_document(doc_name, tasks):
    # Создаём новый документ Word
    # добавить к doc_name .docx
    doc = Document()
    doc.add_heading("Задания", level=1)
    doc.add_heading(f"Уровень сложности {tasks[0][0]["level"]}", level=2)
    for thema in tasks:
        for task in thema:
            doc.add_paragraph(task["task"])
    doc.add_heading("Ответы", level=3)
    for thema in tasks:
        for task in thema:
            doc.add_paragraph(task["correct"])

    doc.save(doc_name)
    messagebox.showinfo("Успех", "файл "+doc_name+" успешно создан")


def task(s, lvl):
    file_names=os.listdir("tasks")
    files = {index: os.path.join("tasks", file) for index, file in enumerate(file_names)}
    #files = {index : os.path.join("tasks", file) for index in range(len(file_names)) for file in file_names}
    tasks = []
    t_create = [i for i, x in enumerate(s) if x == 1]#номера тех которые выбрали
    for elem in t_create:
        file = tc.file_read(files[elem])
        data = tc.prepare_data(file)
        data = tc.filter_level(data, int(lvl))
        data = tc.create_tasks(data, 10)
        tasks.append(data)
    create_word_document(text_entry.get()+'.docx', tasks)


def create_test():
    if text_entry.get() == "":
        messagebox.showerror("Ошибка","Задайте имя для файла!")
        return
    s =[i.get() for i in var_tems]
    if sum(s)==0:
        messagebox.showerror("Ошибка", "Выберите хотя бы одну тему!")
        return
    task(s,lvl=lvl_var.get())


root = tk.Tk()
root.title('glavnoe')
root.geometry("500x500")

lbl = tk.Label(root, text='Добро пожаловать', font='Arial 27')
lbl.place(x=50, y=0, height=50, width=450)
tems = taskcreator.get_task_names("tasks")
var_tems = []
a = 100
for i in range(len(tems)):
    x = tk.BooleanVar()
    var_tems.append(x)
    one = tk.Checkbutton(root, text=tems[i], variable=x)
    a += 30
    one.place(y=a, x=30)
###levl
lvl_var = tk.StringVar(value="1")
lev_1 = tk.Radiobutton(root, value="1", text="1", variable=lvl_var)
lev_1.place(x=210, y=130)
lev_2 = tk.Radiobutton(root, value="2", text="2", variable=lvl_var)
lev_2.place(x=210, y=170)
lev_3 = tk.Radiobutton(root, value="3", text="3", variable=lvl_var)
lev_3.place(x=210, y=210)
###надписи
lbl_tems = tk.Label(root, text='Темы', font='Arial 17')
lbl_tems.place(x=25, y=80)
lbl_tems = tk.Label(root, text='Cложность', font='Arial 17')
lbl_tems.place(x=200, y=80)

text_entry = tk.Entry(root, width=20)
text_entry.place(x=25, y = 300)

btn_add_task = tk.Button(root, bg= "purple", font='Arial 18', text="add task",command=add_window)
btn_add_task.place(x = 270 , y = a+100, height=30, width=150)
btn_Word = tk.Button(root, bg='#4682B4', text='Export Word', font='Arial 18', command=create_test)
btn_Word.place(x=30, y=a + 100, height=30, width=200)


root.mainloop()
