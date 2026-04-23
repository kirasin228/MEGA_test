import tkinter as tk
from tkinter.font import names

from docx import Document
from tkinter import messagebox

bd = []
tas =[]
namess = ""

def create_task_file(name_save):
    global tas
    doc = Document()
    doc.add_paragraph(name_save+"\n")
    for i in tas:
        doc.add_paragraph(i+"\n",)
    doc.save(name_save+".txt")
    messagebox.showinfo("Успех", "файл " + name_save + " успешно создан")

def validat(lis):
    global tas, namess
    for st in lis:
        if st.count(";")!=2:
            messagebox.showerror("Ошибка", "Введите в верном формате")
            return
        a, b, c = st.split(";")
        if a =="" or b =="" or c=="":
            messagebox.showerror("Ошибка", "Введите все значения")
            return
        if c not in ["1","2","3"]:
            messagebox.showerror("Ошибка", "Укажите верный уровень сложности")
            return
    for st in lis:
        if st not in tas:
            tas.append(st)
    create_task_file(namess)


def cr():
    doc = Document()
    doc.save(name.get()+'.docx')
def create_task(win):
    global  bd
    E_tasks = tk.Entry(win, width=14)
    E_tasks.place(x=20,y=win.current_y)
    win.current_y+=25
    bd.append(E_tasks)
def Ok():
    global bd
    l = []
    for i in bd:
        l.append(i.get())
    validat(l)
def create_second_window(parent):
    # Создаем окно уровня Toplevel
    win = tk.Toplevel(parent)
    win.current_y=80
    win.title("Второе окно (функция)")
    win.geometry("300x200")

    # Делаем его зависимым от главного (закроется вместе с ним)
    win.transient(parent)

    # Добавляем элементы
    btn_add_task_win = tk.Button(win,text="Добавить задания",command=lambda: create_task(win))
    btn_add_task_win.place(x=150,y=70)
    name_tsk = tk.Label(win, text="Введите имя задания")
    name_tsk.place(x=20)
    btn_close = tk.Button(win, text="Закрыть", command=win.destroy)
    btn_close.place(x = 150)

    global namess

    name = tk.Entry(win, width=40)
    name.place(x=20,y=40)
    namess = name.get()
    btn_Ok = tk.Button(win, text="Ok",command=Ok)
    btn_Ok.place(x=230)

    # Возвращаем объект окна (не обязательно, но полезно)
    return win
